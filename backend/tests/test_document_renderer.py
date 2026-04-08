"""Tests for DocumentRenderer — DOCX and PDF output correctness."""

import io
from unittest.mock import MagicMock, patch

import pytest
from docx import Document as DocxDocument

from core.models.enums import DocType
from modules.documents.renderer import DocumentRenderer, document_renderer

# ---------------------------------------------------------------------------
# Fixtures & helpers
# ---------------------------------------------------------------------------

SAMPLE_CONTENT: dict = {
    "title": "진단서",
    "diagnosis": "본태성 고혈압(I10)",
    "diagnosis_date": "2026-04-06",
    "clinical_findings": "BP 150/95 mmHg",
    "doctor_opinion": "[의사 소견: ___]",
    "purpose": "보험용",
}

PATIENT_NAME = "홍길동"

_ALL_DOC_TYPES = list(DocType)

# Minimal valid %PDF stub returned by the WeasyPrint mock
_FAKE_PDF = b"%PDF-1.4 fake"


def _mock_weasyprint() -> MagicMock:
    """Return a mock that makes weasyprint.HTML(...).write_pdf() return _FAKE_PDF."""
    wp = MagicMock()
    wp.HTML.return_value.write_pdf.return_value = _FAKE_PDF
    return wp


@pytest.fixture()
def renderer() -> DocumentRenderer:
    return DocumentRenderer()


# ---------------------------------------------------------------------------
# DOCX tests (no system libraries needed)
# ---------------------------------------------------------------------------


def test_render_docx_returns_bytes(renderer: DocumentRenderer) -> None:
    """render_docx must return non-empty bytes."""
    result = renderer.render_docx(DocType.진단서, SAMPLE_CONTENT, PATIENT_NAME)
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_render_docx_contains_title(renderer: DocumentRenderer) -> None:
    """DOCX must contain the document type as a paragraph."""
    result = renderer.render_docx(DocType.진단서, SAMPLE_CONTENT, PATIENT_NAME)
    doc = DocxDocument(io.BytesIO(result))
    texts = [p.text for p in doc.paragraphs]
    assert DocType.진단서.value in texts


def test_render_docx_for_each_doc_type(renderer: DocumentRenderer) -> None:
    """render_docx must produce valid DOCX bytes for every DocType."""
    for doc_type in _ALL_DOC_TYPES:
        result = renderer.render_docx(doc_type, SAMPLE_CONTENT, PATIENT_NAME)
        assert isinstance(result, bytes) and len(result) > 0, f"{doc_type} produced empty DOCX"
        # Must be parseable by python-docx (valid ZIP/DOCX structure)
        DocxDocument(io.BytesIO(result))


def test_render_docx_handles_empty_content(renderer: DocumentRenderer) -> None:
    """render_docx must not crash when content dict is empty."""
    result = renderer.render_docx(DocType.소견서, {}, PATIENT_NAME)
    assert isinstance(result, bytes) and len(result) > 0


# ---------------------------------------------------------------------------
# PDF tests — WeasyPrint mocked to avoid system library requirement
# ---------------------------------------------------------------------------


def test_render_pdf_returns_bytes(renderer: DocumentRenderer) -> None:
    """render_pdf must return non-empty bytes."""
    with patch.dict("sys.modules", {"weasyprint": _mock_weasyprint()}):
        result = renderer.render_pdf(DocType.진단서, SAMPLE_CONTENT, PATIENT_NAME)
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_render_pdf_starts_with_pdf_header(renderer: DocumentRenderer) -> None:
    """PDF bytes must start with the %PDF magic number."""
    with patch.dict("sys.modules", {"weasyprint": _mock_weasyprint()}):
        result = renderer.render_pdf(DocType.진단서, SAMPLE_CONTENT, PATIENT_NAME)
    assert result[:4] == b"%PDF"


def test_render_pdf_for_each_doc_type(renderer: DocumentRenderer) -> None:
    """render_pdf must produce valid PDF bytes for every DocType."""
    for doc_type in _ALL_DOC_TYPES:
        with patch.dict("sys.modules", {"weasyprint": _mock_weasyprint()}):
            result = renderer.render_pdf(doc_type, SAMPLE_CONTENT, PATIENT_NAME)
        assert isinstance(result, bytes) and len(result) > 0, f"{doc_type} produced empty PDF"
        assert result[:4] == b"%PDF", f"{doc_type} PDF does not start with %PDF"


def test_render_pdf_contains_patient_name(renderer: DocumentRenderer) -> None:
    """Jinja2 template must inject patient_name into the rendered HTML.

    We verify this by checking that the HTML passed to WeasyPrint contains
    the patient name, rather than inspecting the opaque PDF binary.
    """
    captured: list[str] = []

    def capture_html(string: str, **_kwargs: object) -> MagicMock:
        captured.append(string)
        mock = MagicMock()
        mock.write_pdf.return_value = _FAKE_PDF
        return mock

    wp = MagicMock()
    wp.HTML.side_effect = capture_html

    with patch.dict("sys.modules", {"weasyprint": wp}):
        renderer.render_pdf(DocType.진단서, SAMPLE_CONTENT, PATIENT_NAME)

    assert captured, "WeasyPrint HTML was never called"
    assert PATIENT_NAME in captured[0], "Patient name not found in rendered HTML"


# ---------------------------------------------------------------------------
# Singleton test
# ---------------------------------------------------------------------------


def test_singleton_instance() -> None:
    """document_renderer module-level singleton must be a DocumentRenderer."""
    assert isinstance(document_renderer, DocumentRenderer)
