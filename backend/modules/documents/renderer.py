"""Document rendering to DOCX and PDF formats.

Uses python-docx for editable Word documents and WeasyPrint for PDF.
Each document type has an HTML template for PDF rendering.
"""

import io
import json
from datetime import date
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from jinja2 import Environment, FileSystemLoader

from core.models.enums import DocType

_TEMPLATE_MAP: dict[DocType, str] = {
    DocType.진단서: "진단서.html",
    DocType.소견서: "소견서.html",
    DocType.의뢰서: "의뢰서.html",
    DocType.확인서: "확인서.html",
    DocType.건강진단서: "건강진단서.html",
}

_SKIP_FIELDS = {"title", "warnings", "patient_info"}


class DocumentRenderer:
    """Renders document content dicts into DOCX or PDF bytes.

    Both render methods accept the same content dict produced by the
    document generation pipeline.  DOCX is built programmatically with
    python-docx; PDF is built via Jinja2 HTML template + WeasyPrint.
    """

    def __init__(self) -> None:
        template_dir = Path(__file__).parent / "templates"
        self._jinja = Environment(
            loader=FileSystemLoader(str(template_dir)),
            autoescape=True,
        )

    def render_docx(
        self,
        doc_type: DocType,
        content: dict,
        patient_name: str,
    ) -> bytes:
        """Build an editable DOCX document from the content dict.

        Args:
            doc_type: Document type enum value used as the title.
            content: Structured content dict from the generation pipeline.
            patient_name: Patient full name for the header line.

        Returns:
            Raw DOCX bytes suitable for HTTP download.
        """
        doc = DocxDocument()

        # Clinic header
        header_para = doc.add_paragraph("보건소")
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.runs[0]
        header_run.bold = True
        header_run.font.size = Pt(10)

        # Document title
        title_para = doc.add_paragraph(doc_type.value)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.bold = True
        title_run.font.size = Pt(16)

        # Patient info
        if patient_name:
            info_para = doc.add_paragraph(f"환자: {patient_name}")
            info_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Horizontal divider via bottom border on an empty paragraph
        doc.add_paragraph("─" * 40)

        # Content fields
        for key, value in content.items():
            if key in _SKIP_FIELDS:
                continue
            heading = doc.add_paragraph()
            heading_run = heading.add_run(key)
            heading_run.bold = True
            heading_run.font.size = Pt(11)

            if isinstance(value, str):
                doc.add_paragraph(value)
            else:
                doc.add_paragraph(json.dumps(value, ensure_ascii=False))

        # Footer
        today = date.today().isoformat()
        doc.add_paragraph("")
        doc.add_paragraph(f"발급일: {today}")
        doc.add_paragraph("담당의: _________________ (서명)")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def render_pdf(
        self,
        doc_type: DocType,
        content: dict,
        patient_name: str,
    ) -> bytes:
        """Render a PDF document via Jinja2 HTML template + WeasyPrint.

        Args:
            doc_type: Document type used to select the HTML template.
            content: Structured content dict from the generation pipeline.
            patient_name: Patient full name injected into the template.

        Returns:
            Raw PDF bytes suitable for HTTP download.
        """
        template_name = self._get_template_name(doc_type)
        template = self._jinja.get_template(template_name)
        today = date.today().isoformat()
        html_str = template.render(
            title=doc_type.value,
            content=content,
            patient_name=patient_name,
            today=today,
        )
        try:
            import weasyprint  # noqa: PLC0415
        except OSError as exc:  # pragma: no cover
            raise RuntimeError(
                "WeasyPrint requires system libraries (pango, gobject). Run: brew install pango"
            ) from exc

        pdf_bytes = weasyprint.HTML(string=html_str).write_pdf()
        return pdf_bytes  # type: ignore[return-value]

    def _get_template_name(self, doc_type: DocType) -> str:
        """Map a DocType to its Jinja2 HTML template filename.

        Args:
            doc_type: Document type enum value.

        Returns:
            Template filename relative to the templates/ directory.
        """
        return _TEMPLATE_MAP.get(doc_type, "default.html")


document_renderer = DocumentRenderer()
